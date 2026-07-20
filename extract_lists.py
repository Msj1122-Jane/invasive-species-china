import docx, pdfplumber, olefile, os, re

# 1. Extract from docx
print('='*60)
print('FILE 1: 口岸重点管控外来物种名录.docx')
print('='*60)
doc = docx.Document(r'D:\OneDrive\Desktop\专题新闻\口岸重点管控外来物种名录.docx')
print(f'Tables: {len(doc.tables)}')
for ti, table in enumerate(doc.tables):
    print(f'--- Table {ti+1} ({len(table.rows)}r x {len(table.columns)}c) ---')
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip()[:50] for cell in row.cells]
        print(' | '.join(cells))
        if ri > 25:
            print(f'... ({len(table.rows)-26} more rows)')
            break
text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
print(f'\nParagraphs text length: {len(text)} chars')
print(text[:2000])

# 2. PDF 2
print('\n' + '='*60)
print('FILE 2: 名单·2.pdf (first 8 pages)')
print('='*60)
try:
    with pdfplumber.open(r'D:\OneDrive\Desktop\专题新闻\名单·2.pdf') as pdf:
        print(f'Total pages: {len(pdf.pages)}')
        for i, page in enumerate(pdf.pages[:8]):
            t = page.extract_text()
            if t:
                print(f'--- Page {i+1} ---')
                print(t[:600])
except Exception as e:
    print(f'PDF2 error: {e}')

# 3. PDF 3
print('\n' + '='*60)
print('FILE 3: 名单3.pdf')
print('='*60)
try:
    with pdfplumber.open(r'D:\OneDrive\Desktop\专题新闻\名单3.pdf') as pdf:
        print(f'Total pages: {len(pdf.pages)}')
        for i, page in enumerate(pdf.pages[:8]):
            t = page.extract_text()
            if t:
                print(f'--- Page {i+1} ---')
                print(t[:600])
except Exception as e:
    print(f'PDF3 error: {e}')

# 4. PDF 4
print('\n' + '='*60)
print('FILE 4: 名单4.pdf (first 8 pages)')
print('='*60)
try:
    with pdfplumber.open(r'D:\OneDrive\Desktop\专题新闻\名单4.pdf') as pdf:
        print(f'Total pages: {len(pdf.pages)}')
        for i, page in enumerate(pdf.pages[:8]):
            t = page.extract_text()
            if t:
                print(f'--- Page {i+1} ---')
                print(t[:600])
except Exception as e:
    print(f'PDF4 error: {e}')

# 5. DOC
print('\n' + '='*60)
print('FILE 5: 1055369430488.doc')
print('='*60)
try:
    ole = olefile.OleFileIO(r'C:\Users\lenovo\Downloads\1055369430488.doc')
    print('OLE streams:', ole.listdir())
    if ole.exists('WordDocument'):
        data = ole.openstream('WordDocument').read()
        # Chinese chars are typically encoded as double-byte
        text_parts = []
        i = 0
        cur = []
        while i < len(data):
            b = data[i]
            if 0x20 <= b < 0x7f or b == 0x0d or b == 0x0a:
                cur.append(chr(b))
            else:
                if len(cur) > 1:
                    t = ''.join(cur).strip()
                    if t and len(t) > 2: text_parts.append(t)
                cur = []
            i += 1
        print(f'Extracted {len(text_parts)} text fragments')
        for p in text_parts[:30]:
            print(p[:120])
except Exception as e:
    print(f'DOC error: {e}')
    import traceback; traceback.print_exc()
